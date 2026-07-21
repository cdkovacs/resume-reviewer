#!/usr/bin/env python3
"""Evaluate a directory of resumes against a skill list and project description
using the Claude API, and export the results as an Excel spreadsheet.

Backends (following the uft-to-flaui-agentic-framework pattern):
    ica        IBM Consulting Advantage (default) — Anthropic-native Messages
               endpoint fronting Anthropic-on-Bedrock. Uses the stock anthropic
               SDK pointed at the /ica base URL.
               Env: ICA_ANTHROPIC_API_KEY (or ICA_API_KEY), ICA_MODEL, ICA_HOST
    anthropic  Direct Claude API. Env: ANTHROPIC_API_KEY, CLAUDE_MODEL

Usage:
    export ICA_ANTHROPIC_API_KEY=<your-key>
    python resume_reviewer.py ./resumes \
        --skills "Python, Kubernetes, Terraform, CI/CD" \
        --project-file project.md \
        --prompt "Weight hands-on infrastructure experience heavily." \
        --output evaluations.xlsx
"""

from __future__ import annotations

import argparse
import base64
import csv
import io
import json
import os
import sys
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from pathlib import Path
from typing import List, Literal, Optional, Type, TypeVar

import anthropic
from pydantic import BaseModel, Field, ValidationError

RESUME_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
DEFAULT_SKILLS_FILE = "skills.csv"
LEGACY_SKILLS_FILE = "skills.txt"  # parses fine as a one-column CSV
DEFAULT_PROJECT_FILE = "project.md"
DEFAULT_STAFF_RATES_FILE = "staff-rates.xlsx"
DEFAULT_RATES_CANDIDATES = ("rates.xlsx", "~/staffing/rates.xlsx")
TEAM_SHAPES = ("low cost", "medium cost", "high cost", "cost not considered")
TEAM_CHARACTERISTICS_INSTRUCTION = (
    "If the project description specifies team characteristics — required roles "
    "(e.g. a project manager), minimum numbers of certain specialists (e.g. a certain "
    "number of scripting specialists), seniority or geographic mixes — treat them as "
    "hard requirements: every team you build must satisfy them wherever the candidate "
    "pool allows, choosing members whose evaluations evidence those roles or "
    "specialties. In composition_notes, state each required characteristic and which "
    "member(s) satisfy it, or explain why it cannot be met with this pool; write "
    "'none specified' if the project description states no team characteristics."
)
DEFAULT_MODEL = "claude-opus-4-8"
DEFAULT_SCREEN_MODEL = "claude-haiku-4-5"
# The /ica root, NOT /ica/v1 — the SDK appends /v1/messages itself.
ICA_DEFAULT_BASE_URL = "https://api.nextgen-beta.ica.ibm.com/ica"


# --------------------------------------------------------------------------
# Structured output schema
# --------------------------------------------------------------------------

class SkillEvaluation(BaseModel):
    skill: str = Field(description="The skill being evaluated, exactly as given in the skills list")
    score: int = Field(description="Score from 0 (no evidence) to 4 (expert, strong direct evidence)")
    evidence: str = Field(description="Specific evidence from the resume supporting the score, or why evidence is lacking")


class ResumeEvaluation(BaseModel):
    candidate_name: str = Field(description="Candidate's full name as it appears on the resume")
    overall_score: int = Field(description="Overall fit for the project, 0 (poor fit) to 4 (excellent fit)")
    recommendation: Literal["strong yes", "yes", "maybe", "no"]
    summary: str = Field(description="2-4 sentence assessment of the candidate's fit for the project")
    strengths: List[str] = Field(description="Top strengths relevant to the project")
    gaps: List[str] = Field(description="Notable gaps or risks relative to the project needs")
    skill_evaluations: List[SkillEvaluation] = Field(
        description="One entry per skill in the skills list, in the same order"
    )
    email: str = Field(
        description="Candidate's email address as it appears on the resume; "
        "empty string if none is present"
    )
    candidate_feedback: str = Field(
        description="A message to the candidate in a respectful, professional, "
        "positive tone, formatted like an email: start with 'Dear <name>,' followed "
        "by a blank line (two newlines); then the body paragraphs; then a blank line "
        "(two newlines) before a final closing line that begins 'We appreciate your "
        "time'. The body must: (1) tell them they are being considered for a seat on "
        "a project; (2) briefly describe what we are looking for in a candidate, in "
        "general terms without disclosing confidential client details; (3) highlight "
        "the skill areas where their background matched well; and (4) ask whether "
        "they have experience in the skill areas where the resume showed little or "
        "no evidence, inviting them to share it. Write in complete, flowing "
        "sentences: no dashes, em dashes, hyphens-as-punctuation, or bullet lists. "
        "Never mention numeric scores, rankings, or other candidates."
    )


@dataclass(frozen=True)
class Skill:
    name: str
    weight: float = 1.0


def parse_skills_csv(text: str, source: str) -> List[Skill]:
    """Parse `skill[,weight]` lines; weight defaults to 1.0. Blank lines and
    #-comments are skipped. A plain one-skill-per-line file parses unchanged."""
    skills: List[Skill] = []
    for row in csv.reader(io.StringIO(text)):
        if not row or not row[0].strip() or row[0].lstrip().startswith("#"):
            continue
        name = row[0].strip()
        weight = 1.0
        if len(row) > 1 and row[1].strip():
            try:
                weight = float(row[1])
            except ValueError:
                sys.exit(f"error: invalid weight {row[1].strip()!r} for skill {name!r} in {source}")
            if weight <= 0:
                sys.exit(f"error: weight for skill {name!r} in {source} must be > 0")
        skills.append(Skill(name=name, weight=weight))
    return skills


class TeamSelection(BaseModel):
    selected_candidates: List[str] = Field(
        description="Names of exactly the requested number of candidates, copied verbatim "
        "from the provided candidate list"
    )
    rationale: str = Field(
        description="Why this mix forms the strongest overall team: skill coverage, "
        "complementary profiles, and any tradeoffs made"
    )
    composition_notes: str = Field(
        description="How each team characteristic required by the project description "
        "(required roles, specialist counts, mixes) is satisfied by this team, or why "
        "it cannot be with this candidate pool; 'none specified' if the project "
        "description states no team characteristics"
    )


class TeamShapeSelection(BaseModel):
    shape: Literal["low cost", "medium cost", "high cost", "cost not considered"]
    members: List[str] = Field(
        description="Names of exactly the requested number of candidates for this shape, "
        "copied verbatim from the provided list"
    )
    rationale: str = Field(description="Why this mix fits the shape's cost/skill objective")
    composition_notes: str = Field(
        description="How each team characteristic required by the project description "
        "(required roles, specialist counts, mixes) is satisfied by this team, or why "
        "it cannot be with this candidate pool; 'none specified' if the project "
        "description states no team characteristics"
    )


class TeamShapesSelection(BaseModel):
    shapes: List[TeamShapeSelection] = Field(
        description="Exactly four entries, one per shape: low cost, medium cost, "
        "high cost, cost not considered"
    )


class ScreenResult(BaseModel):
    candidate_name: str = Field(description="Candidate's full name as it appears on the resume")
    overall_score: int = Field(description="Overall fit for the project, 0 (poor fit) to 4 (excellent fit)")
    recommendation: Literal["strong yes", "yes", "maybe", "no"]
    summary: str = Field(description="1-2 sentence assessment of the candidate's fit")


class CandidateName(BaseModel):
    candidate_name: str = Field(description="The candidate's full name as it appears on the resume")


@dataclass
class Result:
    file: Path
    evaluation: Optional[ResumeEvaluation] = None
    error: Optional[str] = None
    stage: str = "full"  # "full" or "screened out"


# --------------------------------------------------------------------------
# Resume loading
# --------------------------------------------------------------------------

def extract_docx_text(path: Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def resume_content_block(path: Path) -> dict:
    """Build the message content block carrying the resume itself."""
    if path.suffix.lower() == ".pdf":
        data = base64.standard_b64encode(path.read_bytes()).decode("utf-8")
        return {
            "type": "document",
            "source": {"type": "base64", "media_type": "application/pdf", "data": data},
        }
    if path.suffix.lower() == ".docx":
        text = extract_docx_text(path)
    else:
        text = path.read_text(encoding="utf-8", errors="replace")
    if not text.strip():
        raise ValueError("resume file is empty after text extraction")
    return {"type": "text", "text": f"<resume>\n{text}\n</resume>"}


# --------------------------------------------------------------------------
# Staff rates
# --------------------------------------------------------------------------

@dataclass
class StaffInfo:
    name: str
    geo: str = ""
    line: str = ""
    band: str = ""
    rate: Optional[float] = None
    currency: str = ""


@dataclass
class RatesTable:
    rates: dict  # (geo, line, band) -> (rate, currency), keys normalized
    rows: List[tuple]  # original (country, line, band, rate, currency) rows, in file order
    geos: List[object]
    lines: List[object]
    bands: List[object]
    currencies: List[object]


def _norm(value) -> str:
    return str(value).strip().upper() if value is not None else ""


def load_rates(path: Path) -> RatesTable:
    """Rates table with header row Country, Line, Band, Rate, Currency.
    Distinct original values are kept (in file order) for dropdown lists."""
    from openpyxl import load_workbook

    ws = load_workbook(path, data_only=True).active
    assert ws is not None
    rates: dict = {}
    rows: List[tuple] = []
    seen: dict = {"geo": {}, "line": {}, "band": {}, "currency": {}}
    for row in ws.iter_rows(min_row=2, values_only=True):
        country, line, band, rate, currency = (list(row) + [None] * 5)[:5]
        if country is None or line is None or band is None or rate is None:
            continue
        try:
            rate_value = float(rate)  # type: ignore[arg-type]
        except (TypeError, ValueError):
            continue
        rates[(_norm(country), _norm(line), _norm(band))] = (rate_value, str(currency or ""))
        rows.append((country, line, band, rate, currency))
        seen["geo"].setdefault(_norm(country), country)
        seen["line"].setdefault(_norm(line), line)
        seen["band"].setdefault(_norm(band), band)
        if currency:
            seen["currency"].setdefault(_norm(currency), currency)
    return RatesTable(
        rates=rates,
        rows=rows,
        geos=list(seen["geo"].values()),
        lines=list(seen["line"].values()),
        bands=list(seen["band"].values()),
        currencies=list(seen["currency"].values()),
    )


def apply_staff_dropdowns(wb, ws, rates: RatesTable) -> None:
    """Cascading data-validation dropdowns for the fill-in fields.

    Geo and Currency are flat lists. Line is filtered by the chosen Geo, and
    Band by the chosen Geo+Line, using OFFSET/MATCH/COUNTIF over sorted helper
    tables on a hidden Lookups sheet (the INDIRECT/named-range trick can't
    work here: line names contain spaces, hyphens, and '&', which are invalid
    in defined names). Helper rows are sorted so each parent's children are
    contiguous — required for the OFFSET window to be correct."""
    from openpyxl.worksheet.datavalidation import DataValidation

    if "Lookups" in wb.sheetnames:
        del wb["Lookups"]
    lk = wb.create_sheet("Lookups")
    lk.sheet_state = "hidden"

    # Flat lists: Geo (A), Currency (B)
    lk["A1"] = "Geo"
    for i, value in enumerate(rates.geos, start=2):
        lk[f"A{i}"] = value
    lk["B1"] = "Currency"
    for i, value in enumerate(rates.currencies, start=2):
        lk[f"B{i}"] = value

    # Cascade helpers: distinct (Geo, Line) pairs sorted by Geo (D:E), and
    # distinct (Geo|Line key, Band) sorted by key (G:H).
    pair_map: dict = {}
    triple_map: dict = {}
    for country, line, band, _rate, _currency in rates.rows:
        pair_map.setdefault((_norm(country), _norm(line)), (country, line))
        triple_map.setdefault((_norm(country), _norm(line), _norm(band)), (country, line, band))
    pairs = [pair_map[k] for k in sorted(pair_map)]
    triples = [triple_map[k] for k in sorted(triple_map)]

    lk["D1"], lk["E1"] = "PairGeo", "PairLine"
    for i, (country, line) in enumerate(pairs, start=2):
        lk[f"D{i}"], lk[f"E{i}"] = country, line
    lk["G1"], lk["H1"] = "GeoLineKey", "Band"
    for i, (country, line, band) in enumerate(triples, start=2):
        lk[f"G{i}"], lk[f"H{i}"] = f"{country}|{line}", band

    ws.data_validations.dataValidation = []  # rebuilt fresh on every sync
    last_row = max(ws.max_row + 200, 1000)
    np, nt = len(pairs) + 1, len(triples) + 1
    validations = [
        # Geo: flat list
        ("B", f"Lookups!$A$2:$A${len(rates.geos) + 1}"),
        # Line: rows of the pair table whose Geo matches $B2
        ("C", f"OFFSET(Lookups!$E$1,MATCH($B2,Lookups!$D$2:$D${np},0),0,"
              f"COUNTIF(Lookups!$D$2:$D${np},$B2),1)"),
        # Band: rows of the triple table whose Geo|Line key matches $B2|$C2
        ("D", f'OFFSET(Lookups!$H$1,MATCH($B2&"|"&$C2,Lookups!$G$2:$G${nt},0),0,'
              f'COUNTIF(Lookups!$G$2:$G${nt},$B2&"|"&$C2),1)'),
        # Currency: flat list
        ("F", f"Lookups!$B$2:$B${len(rates.currencies) + 1}"),
    ]
    for target_col, formula in validations:
        dv = DataValidation(type="list", formula1=formula, allow_blank=True)
        ws.add_data_validation(dv)
        dv.add(f"{target_col}2:{target_col}{last_row}")


def write_rates_sheet(wb, rates: RatesTable) -> None:
    """Copy the rates table into a 'Rates' tab so the workbook is
    self-contained. Column D holds a hidden Geo|Line|Band key that the
    Staff Rates sheet's VLOOKUP formulas resolve against."""
    from openpyxl.styles import Font, PatternFill

    if "Rates" in wb.sheetnames:
        del wb["Rates"]
    rs = wb.create_sheet("Rates")
    rs.append(["Country", "Line", "Band", "Key", "Rate", "Currency"])
    for cell in rs[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="D9E1F2")
    rs.freeze_panes = "A2"
    for i, (country, line, band, rate, currency) in enumerate(rates.rows, start=2):
        # The key is an Excel formula (not a Python string) so its text
        # coercion of numeric bands matches the lookup side exactly.
        rs.append([country, line, band, f'=A{i}&"|"&B{i}&"|"&C{i}', rate, currency])
    rs.column_dimensions["D"].hidden = True
    for col, width in zip("ABCEF", [10, 42, 8, 10, 10]):
        rs.column_dimensions[col].width = width


def sync_staff_rates(
    staff_path: Path,
    candidate_names: List[str],
    rates: Optional[RatesTable],
) -> dict:
    """Create/update staff-rates.xlsx and return {normalized name: StaffInfo}.

    Existing rows are never modified except to fill in Rate/Currency when
    Geo/Line/Band are populated and a rates table is available. Candidates not
    yet present are appended with blank Geo/Line/Band for the staffing team to
    fill in.
    """
    from openpyxl import Workbook, load_workbook

    headers = ["Name", "Geo", "Line", "Band", "Rate", "Currency"]
    if staff_path.is_file():
        wb = load_workbook(staff_path)
        ws = wb.active
        assert ws is not None
        if [c.value for c in ws[1]][: len(headers)] != headers:
            sys.exit(f"error: {staff_path} does not have the expected header row {headers}")
    else:
        wb = Workbook()
        ws = wb.active
        assert ws is not None
        ws.title = "Staff Rates"
        ws.append(headers)
        from openpyxl.styles import Font, PatternFill

        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill("solid", fgColor="D9E1F2")
        ws.freeze_panes = "A2"
        for col, width in zip("ABCDEF", [28, 10, 30, 8, 10, 10]):
            ws.column_dimensions[col].width = width

    # In-workbook lookup: Rate/Currency cells become VLOOKUP formulas against
    # the Rates tab, so they populate live in Excel as soon as Geo/Line/Band
    # are chosen — no re-run needed.
    lookup_range = f"Rates!$D$2:$F${len(rates.rows) + 1}" if rates is not None else None

    def set_rate_formulas(row_idx: int) -> None:
        key = f'B{row_idx}&"|"&C{row_idx}&"|"&D{row_idx}'
        blank_guard = f'OR(B{row_idx}="",C{row_idx}="",D{row_idx}="")'
        ws.cell(row=row_idx, column=5).value = (
            f'=IF({blank_guard},"",IFERROR(VLOOKUP({key},{lookup_range},2,FALSE),"no match"))'
        )
        ws.cell(row=row_idx, column=6).value = (
            f'=IF({blank_guard},"",IFERROR(VLOOKUP({key},{lookup_range},3,FALSE),""))'
        )

    staff: dict = {}
    unmatched: List[str] = []
    for row in ws.iter_rows(min_row=2):
        name = row[0].value
        if not name or not str(name).strip():
            continue
        raw_rate, raw_currency = row[4].value, row[5].value
        info = StaffInfo(
            name=str(name).strip(),
            geo=str(row[1].value or "").strip(),
            line=str(row[2].value or "").strip(),
            band=str(row[3].value or "").strip(),
            # Cells may now hold formulas; only literal values are read here.
            rate=float(raw_rate) if isinstance(raw_rate, (int, float)) else None,
            currency=raw_currency.strip()
            if isinstance(raw_currency, str) and not raw_currency.startswith("=")
            else "",
        )
        if rates is not None:
            # Python-side rate for team shapes comes straight from the rates
            # table — independent of whether Excel has recalculated formulas.
            if info.geo and info.line and info.band:
                hit = rates.rates.get((_norm(info.geo), _norm(info.line), _norm(info.band)))
                if hit:
                    info.rate, info.currency = hit
                else:
                    info.rate, info.currency = None, ""
                    unmatched.append(f"{info.name} ({info.geo}/{info.line}/{info.band})")
            set_rate_formulas(row[0].row)
        staff[info.name.lower()] = info

    added = 0
    for name in candidate_names:
        if name.strip().lower() not in staff:
            ws.append([name, "", "", "", None, ""])
            if rates is not None:
                set_rate_formulas(ws.max_row)
            staff[name.strip().lower()] = StaffInfo(name=name)
            added += 1

    if rates is not None:
        write_rates_sheet(wb, rates)
        apply_staff_dropdowns(wb, ws, rates)
    wb.save(staff_path)
    if added:
        print(f"staff-rates: added {added} new candidate(s) to {staff_path} — "
              "fill in Geo/Line/Band and re-run to compute rates")
    if unmatched:
        print("warning: no rate found in the rates table for: " + "; ".join(unmatched))
    return staff


# --------------------------------------------------------------------------
# Backends
# --------------------------------------------------------------------------

def make_client(provider: str, host: Optional[str] = None) -> anthropic.Anthropic:
    """Build an Anthropic-SDK client for the chosen provider.

    ICA exposes an Anthropic-native Messages endpoint (/ica/v1/messages), so the
    stock SDK works unchanged with a base_url override — same approach as
    uft-to-flaui-agentic-framework's ICABackend.
    """
    if provider == "ica":
        key = os.environ.get("ICA_ANTHROPIC_API_KEY") or os.environ.get("ICA_API_KEY")
        if not key:
            sys.exit("error: ICA_ANTHROPIC_API_KEY (or ICA_API_KEY) is not set")
        base_url = (host or os.environ.get("ICA_HOST", ICA_DEFAULT_BASE_URL)).rstrip("/")
        return anthropic.Anthropic(api_key=key, base_url=base_url)
    if provider == "anthropic":
        if not os.environ.get("ANTHROPIC_API_KEY"):
            sys.exit("error: ANTHROPIC_API_KEY is not set")
        return anthropic.Anthropic()
    sys.exit(f"error: unknown provider {provider!r} (choose: ica, anthropic)")


def resolve_model(provider: str, explicit: Optional[str]) -> str:
    """--model flag > provider-specific env var > default."""
    env_key = "ICA_MODEL" if provider == "ica" else "CLAUDE_MODEL"
    return explicit or os.environ.get(env_key) or DEFAULT_MODEL


# --------------------------------------------------------------------------
# Evaluation
# --------------------------------------------------------------------------

def build_system_prompt(prompt: str, project: str, skills: List[Skill]) -> str:
    skills_list = "\n".join(f"- {s.name} (weight: {s.weight:g})" for s in skills)
    sections = [
        "You are an experienced technical recruiter and hiring manager. "
        "Evaluate the resume you are given against the project description and skills below. "
        "Base every score strictly on evidence in the resume; do not give the benefit of the doubt "
        "for skills that are not demonstrated. Score each skill from 0 (no evidence) to 4 "
        "(expert with strong direct evidence). Provide one skill_evaluations entry per listed "
        "skill, using the exact skill names given, in the same order. "
        "Each skill's weight indicates its relative importance to this project: individual "
        "skill scores stay evidence-based and unweighted, but weights should inform the "
        "overall fit score, recommendation, and any team selection.",
        f"<project_description>\n{project}\n</project_description>",
        f"<skills_to_evaluate>\n{skills_list}\n</skills_to_evaluate>",
    ]
    if prompt:
        sections.append(f"<additional_evaluation_instructions>\n{prompt}\n</additional_evaluation_instructions>")
    return "\n\n".join(sections)


T = TypeVar("T", bound=BaseModel)


def call_structured(
    client: anthropic.Anthropic,
    model: str,
    system_prompt: str,
    user_content: list,
    schema_model: Type[T],
    tool_name: str,
    max_tokens: int,
) -> T:
    """Tool-forced structured output, validated with Pydantic.

    Forcing a single tool whose input_schema is the result schema works on both
    the direct Claude API and ICA's Bedrock-fronted Messages endpoint (the
    surface uft-to-flaui-agentic-framework verified live), unlike
    output_config.format which is unverified through the proxy. Forced
    tool_choice also means no thinking parameter. One retry on validation
    failure, feeding the error back as a tool_result.
    """
    tool = {
        "name": tool_name,
        "description": "Record the evaluation result, strictly conforming to the schema.",
        "input_schema": schema_model.model_json_schema(),
    }
    messages: list = [{"role": "user", "content": user_content}]
    last_err: Optional[ValidationError] = None
    for attempt in range(2):
        response = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            system=[
                {
                    "type": "text",
                    "text": system_prompt,
                    # Shared across every resume in the run — cache it.
                    "cache_control": {"type": "ephemeral"},
                }
            ],
            tools=[tool],
            tool_choice={"type": "tool", "name": tool_name},
            messages=messages,
        )
        if response.stop_reason == "refusal":
            raise RuntimeError("the model declined to evaluate this resume (stop_reason: refusal)")
        if response.stop_reason == "max_tokens":
            raise RuntimeError("evaluation was truncated (stop_reason: max_tokens)")
        block = next(
            (b for b in response.content if b.type == "tool_use" and b.name == tool_name),
            None,
        )
        if block is None:
            raise RuntimeError("model did not return the expected tool_use block")
        try:
            return schema_model.model_validate(block.input)
        except ValidationError as exc:
            last_err = exc
            if attempt == 0:
                messages = messages + [
                    {"role": "assistant", "content": response.content},
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "tool_result",
                                "tool_use_id": block.id,
                                "content": f"Arguments failed validation: {exc}. "
                                f"Call {tool_name} again with corrected arguments.",
                                "is_error": True,
                            }
                        ],
                    },
                ]
    raise RuntimeError(f"model output failed schema validation after retry: {last_err}")


def resume_user_content(path: Path, instruction: str) -> list:
    """Resume block plus instruction, with the file name as extra context —
    file names often carry the candidate's name (e.g. Jane_Smith_Resume.pdf)."""
    return [
        resume_content_block(path),
        {
            "type": "text",
            "text": f'The resume file is named "{path.name}" — use it as additional '
            f"context for the candidate's name. {instruction}",
        },
    ]


def evaluate_resume(
    client: anthropic.Anthropic,
    model: str,
    system_prompt: str,
    path: Path,
) -> ResumeEvaluation:
    return call_structured(
        client,
        model,
        system_prompt,
        resume_user_content(path, "Evaluate this resume."),
        ResumeEvaluation,
        tool_name="record_evaluation",
        max_tokens=16000,
    )


def screen_resume(
    client: anthropic.Anthropic,
    model: str,
    system_prompt: str,
    path: Path,
) -> ScreenResult:
    """Cheap first-pass screen with a lower-cost model (no per-skill breakdown)."""
    return call_structured(
        client,
        model,
        system_prompt,
        resume_user_content(
            path,
            "Quickly screen this resume: give only the candidate name, "
            "an overall fit score, a recommendation, and a 1-2 sentence summary.",
        ),
        ScreenResult,
        tool_name="record_screen_result",
        max_tokens=2000,
    )


def extract_name(client: anthropic.Anthropic, model: str, path: Path) -> str:
    """Name-only extraction with a cheap model — used to bootstrap
    staff-rates.xlsx without running full skill evaluations."""
    result = call_structured(
        client,
        model,
        "Extract the candidate's full name from the resume you are given.",
        resume_user_content(path, "Extract the candidate's name."),
        CandidateName,
        tool_name="record_name",
        max_tokens=500,
    )
    return result.candidate_name


def find_rates_table(rates_arg: Optional[str]) -> Optional[RatesTable]:
    """--rates path, else the first default candidate that exists."""
    candidates = [rates_arg] if rates_arg else list(DEFAULT_RATES_CANDIDATES)
    for cand in candidates:
        p = Path(cand).expanduser()
        if p.is_file():
            table = load_rates(p)
            print(f"Loaded {len(table.rates)} rate entries from {p}")
            return table
    if rates_arg:
        sys.exit(f"error: rates file {rates_arg} not found")
    print("note: no rates table found (looked for " + ", ".join(candidates) + ")")
    return None


def select_team(
    client: anthropic.Anthropic,
    model: str,
    system_prompt: str,
    pool: List[tuple],
    size: int,
) -> TeamSelection:
    """Ask the model to pick the strongest complementary team from the pool.

    Reuses the same cached system prompt (project + skills) as the evaluations.
    """
    candidates = []
    for _, ev in pool:
        candidates.append(
            {
                "name": ev.candidate_name,
                "overall_score": ev.overall_score,
                "recommendation": ev.recommendation,
                "skill_scores": {se.skill: se.score for se in ev.skill_evaluations},
                "strengths": ev.strengths,
                "gaps": ev.gaps,
                "summary": ev.summary,
            }
        )
    user_text = (
        f"Below are structured evaluations of {len(pool)} candidates for this project.\n\n"
        f"{json.dumps(candidates, indent=1)}\n\n"
        f"Select the {size} candidates who together form the strongest overall team for the "
        "project. Optimize the mix, not just individual scores: every skill area should be "
        "covered by at least one strong member, weight the project's critical initiatives "
        "most heavily, and prefer complementary profiles over redundant ones. "
        + TEAM_CHARACTERISTICS_INSTRUCTION
        + f" Return exactly {size} names, copied verbatim from the list above."
    )
    return call_structured(
        client,
        model,
        system_prompt,
        [{"type": "text", "text": user_text}],
        TeamSelection,
        tool_name="record_team_selection",
        max_tokens=8000,
    )


@dataclass
class ShapeResult:
    shape: str
    members: List[tuple]  # [(Result, ResumeEvaluation)]
    rationale: str
    fit_score: float
    avg_rate: float
    currency: str


def team_fit_score(members: List[tuple], skills: List[Skill]) -> float:
    """Weighted skill coverage: each skill counts as the best member's score."""
    total = 0.0
    for s in skills:
        key = s.name.strip().lower()
        best = max(
            (se.score for _, ev in members for se in ev.skill_evaluations
             if se.skill.strip().lower() == key),
            default=0,
        )
        total += best * s.weight
    return round(total / sum(s.weight for s in skills), 2)


def shape_metrics(shape: str, members: List[tuple], rationale: str,
                  skills: List[Skill], staff: dict) -> ShapeResult:
    rates = [staff[ev.candidate_name.strip().lower()].rate for _, ev in members]
    currencies = sorted({staff[ev.candidate_name.strip().lower()].currency for _, ev in members})
    return ShapeResult(
        shape=shape,
        members=members,
        rationale=rationale,
        fit_score=team_fit_score(members, skills),
        avg_rate=round(sum(rates) / len(rates), 2),
        currency=currencies[0] if len(currencies) == 1 else "MIXED: " + ", ".join(currencies),
    )


def greedy_best_fit(pool: List[tuple], skills: List[Skill], size: int) -> List[tuple]:
    """Fallback: iteratively add the candidate that raises team fit the most."""
    team: List[tuple] = []
    remaining = list(pool)
    while len(team) < size and remaining:
        best = max(remaining, key=lambda entry: (team_fit_score(team + [entry], skills),
                                                 entry[1].overall_score))
        team.append(best)
        remaining.remove(best)
    return team


def select_team_shapes(
    client: anthropic.Anthropic,
    model: str,
    system_prompt: str,
    pool: List[tuple],
    staff: dict,
    size: int,
) -> TeamShapesSelection:
    """One call: pick four teams of `size` optimizing different cost postures."""
    candidates = []
    for _, ev in pool:
        info = staff[ev.candidate_name.strip().lower()]
        candidates.append(
            {
                "name": ev.candidate_name,
                "overall_score": ev.overall_score,
                "skill_scores": {se.skill: se.score for se in ev.skill_evaluations},
                "strengths": ev.strengths,
                "gaps": ev.gaps,
                "cost_rate": info.rate,
                "currency": info.currency,
            }
        )
    user_text = (
        f"Below are structured evaluations of {len(pool)} candidates, each with an hourly "
        f"cost rate.\n\n{json.dumps(candidates, indent=1)}\n\n"
        f"Build FOUR teams of exactly {size} members each — every team must have a decent "
        "skill fit (each skill area covered by at least one capable member wherever the pool "
        "allows, weighted toward the project's critical initiatives), differing in cost "
        "posture:\n"
        "1. 'low cost' — minimize average cost rate while keeping an acceptable skill fit.\n"
        "2. 'medium cost' — balance cost and skill fit.\n"
        "3. 'high cost' — favor premium, senior talent where it buys real skill.\n"
        "4. 'cost not considered' — the strongest possible skill fit, ignoring cost.\n"
        + TEAM_CHARACTERISTICS_INSTRUCTION
        + " These characteristics apply to ALL four shapes — cost posture never overrides "
        "them. Candidates may appear in multiple teams. Return names copied verbatim."
    )
    return call_structured(
        client,
        model,
        system_prompt,
        [{"type": "text", "text": user_text}],
        TeamShapesSelection,
        tool_name="record_team_shapes",
        max_tokens=16000,
    )


# --------------------------------------------------------------------------
# Excel export
# --------------------------------------------------------------------------

RECOMMENDATION_FILLS = {
    "strong yes": "C6EFCE",
    "yes": "E2EFDA",
    "maybe": "FFF2CC",
    "no": "F8CBAD",
}


def write_workbook(
    results: List[Result],
    skills: List[Skill],
    output: Path,
    team: Optional[tuple] = None,  # (members: List[(Result, ResumeEvaluation)], rationale: str)
    shapes: Optional[List[ShapeResult]] = None,
    staff: Optional[dict] = None,
) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    header_font = Font(bold=True)
    header_fill = PatternFill("solid", fgColor="D9E1F2")
    wrap = Alignment(wrap_text=True, vertical="top")

    def style_header(ws) -> None:
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
        ws.freeze_panes = "A2"

    def skill_label(s: Skill) -> str:
        return f"{s.name} (0-4)" if s.weight == 1.0 else f"{s.name} (0-4, w={s.weight:g})"

    def weighted_avg(ev: ResumeEvaluation) -> object:
        scores_by_skill = {se.skill.strip().lower(): se.score for se in ev.skill_evaluations}
        pairs = [
            (scores_by_skill[s.name.strip().lower()], s.weight)
            for s in skills
            if s.name.strip().lower() in scores_by_skill
        ]
        if not pairs:
            return ""
        return round(sum(score * w for score, w in pairs) / sum(w for _, w in pairs), 2)

    # --- Summary sheet -----------------------------------------------------
    ws = wb.active
    assert ws is not None
    ws.title = "Summary"
    headers = ["Candidate", "File", "Overall (0-4)", "Weighted avg (0-4)", "Recommendation", "Stage"]
    headers += [skill_label(s) for s in skills]
    headers += ["Strengths", "Gaps", "Summary"]
    ws.append(headers)

    ok = [(r, r.evaluation) for r in results if r.evaluation is not None]
    ok.sort(key=lambda pair: pair[1].overall_score, reverse=True)

    for r, ev in ok:
        scores_by_skill = {se.skill.strip().lower(): se.score for se in ev.skill_evaluations}
        row = [ev.candidate_name, r.file.name, ev.overall_score, weighted_avg(ev),
               ev.recommendation, r.stage]
        row += [scores_by_skill.get(s.name.strip().lower(), "") for s in skills]
        row += ["\n".join(ev.strengths), "\n".join(ev.gaps), ev.summary]
        ws.append(row)
        rec_cell = ws.cell(row=ws.max_row, column=5)
        fill = RECOMMENDATION_FILLS.get(ev.recommendation)
        if fill:
            rec_cell.fill = PatternFill("solid", fgColor=fill)
        for col in (len(headers) - 2, len(headers) - 1, len(headers)):
            ws.cell(row=ws.max_row, column=col).alignment = wrap

    style_header(ws)
    widths = [24, 24, 12, 14, 15, 13] + [14] * len(skills) + [45, 45, 60]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # --- Team Shapes sheet (second tab, when rates enabled the four shapes) --
    if shapes:
        ws_s = wb.create_sheet("Team Shapes", 1)
        ws_s.append(["Shape", "Skill fit (0-4)", "Avg cost rate", "Currency", "Members"])
        for sh in shapes:
            ws_s.append([
                sh.shape, sh.fit_score, sh.avg_rate, sh.currency,
                ", ".join(ev.candidate_name for _, ev in sh.members),
            ])
            ws_s.cell(row=ws_s.max_row, column=5).alignment = wrap
        style_header(ws_s)

        block_headers = ["Candidate", "Overall (0-4)", "Rate", "Currency"] + [
            skill_label(s) for s in skills
        ]
        for sh in shapes:
            ws_s.append([])
            ws_s.append([f"{sh.shape}  —  skill fit {sh.fit_score}, avg rate "
                         f"{sh.avg_rate} {sh.currency}"])
            title_cell = ws_s.cell(row=ws_s.max_row, column=1)
            title_cell.font = header_font
            title_cell.fill = header_fill
            ws_s.merge_cells(start_row=ws_s.max_row, start_column=1,
                             end_row=ws_s.max_row, end_column=len(block_headers))
            ws_s.append(block_headers)
            for cell in ws_s[ws_s.max_row]:
                cell.font = header_font
            for _, ev in sh.members:
                info = (staff or {}).get(ev.candidate_name.strip().lower())
                scores_by_skill = {se.skill.strip().lower(): se.score
                                   for se in ev.skill_evaluations}
                ws_s.append(
                    [ev.candidate_name, ev.overall_score,
                     info.rate if info else "", info.currency if info else ""]
                    + [scores_by_skill.get(s.name.strip().lower(), "") for s in skills]
                )
            totals = []
            for s in skills:
                key = s.name.strip().lower()
                totals.append(sum(
                    se.score for _, ev in sh.members for se in ev.skill_evaluations
                    if se.skill.strip().lower() == key
                ))
            ws_s.append(["Team skill total", "", "", ""] + totals)
            for cell in ws_s[ws_s.max_row]:
                cell.font = header_font
            ws_s.append(["Rationale", sh.rationale])
            rat = ws_s.cell(row=ws_s.max_row, column=2)
            rat.alignment = wrap
            ws_s.merge_cells(start_row=ws_s.max_row, start_column=2,
                             end_row=ws_s.max_row, end_column=len(block_headers))
        ws_s.column_dimensions["A"].width = 26
        for col, width in zip("BCD", [13, 12, 10]):
            ws_s.column_dimensions[col].width = width
        ws_s.column_dimensions["E"].width = 40
        for i in range(5, len(block_headers) + 1):
            ws_s.column_dimensions[get_column_letter(i)].width = 14

    # --- Team sheet (second tab, when a target team size was given) --------
    if team is not None:
        members, rationale = team
        ws_t = wb.create_sheet("Team", 1)
        t_headers = ["Candidate", "Overall (0-4)"] + [skill_label(s) for s in skills]
        ws_t.append(t_headers)
        for _, ev in members:
            scores_by_skill = {se.skill.strip().lower(): se.score for se in ev.skill_evaluations}
            ws_t.append(
                [ev.candidate_name, ev.overall_score]
                + [scores_by_skill.get(s.name.strip().lower(), "") for s in skills]
            )
        totals = []
        for s in skills:
            key = s.name.strip().lower()
            totals.append(
                sum(
                    se.score
                    for _, ev in members
                    for se in ev.skill_evaluations
                    if se.skill.strip().lower() == key
                )
            )
        ws_t.append(["Team skill total", ""] + totals)
        for cell in ws_t[ws_t.max_row]:
            cell.font = header_font
            cell.fill = header_fill
        ws_t.append([])
        ws_t.append(["Selection rationale", rationale])
        rat_cell = ws_t.cell(row=ws_t.max_row, column=2)
        rat_cell.alignment = wrap
        ws_t.merge_cells(
            start_row=ws_t.max_row, start_column=2,
            end_row=ws_t.max_row, end_column=len(t_headers),
        )
        style_header(ws_t)
        ws_t.column_dimensions["A"].width = 24
        ws_t.column_dimensions["B"].width = 12
        for i in range(3, len(t_headers) + 1):
            ws_t.column_dimensions[get_column_letter(i)].width = 14

    # --- Details sheet -----------------------------------------------------
    ws_d = wb.create_sheet("Skill Details")
    ws_d.append(["Candidate", "Skill", "Score (0-4)", "Evidence"])
    for r, ev in ok:
        for se in ev.skill_evaluations:
            ws_d.append([ev.candidate_name, se.skill, se.score, se.evidence])
            ws_d.cell(row=ws_d.max_row, column=4).alignment = wrap
    style_header(ws_d)
    for col, w in zip("ABCD", [24, 24, 12, 90]):
        ws_d.column_dimensions[col].width = w

    # --- Ranking sheet: force-ranked bar chart of weighted averages ---------
    ranked = sorted(
        (
            (r, ev, weighted_avg(ev))
            for r, ev in ok
            if isinstance(weighted_avg(ev), (int, float))
        ),
        key=lambda t: (t[2], t[1].overall_score, t[1].candidate_name),
        reverse=True,
    )
    if ranked:
        from openpyxl.chart import BarChart, Reference
        from openpyxl.chart.label import DataLabelList

        ws_r = wb.create_sheet("Ranking")
        ws_r.append(["Rank", "Candidate", "Weighted avg (0-4)"])
        for rank, (_, ev, wa) in enumerate(ranked, start=1):
            ws_r.append([rank, ev.candidate_name, wa])
        style_header(ws_r)
        ws_r.column_dimensions["B"].width = 28
        ws_r.column_dimensions["C"].width = 18

        chart = BarChart()
        chart.type = "bar"  # horizontal bars — readable candidate names
        chart.title = "Force ranking — weighted average (0-4)"
        data = Reference(ws_r, min_col=3, min_row=1, max_row=len(ranked) + 1)
        cats = Reference(ws_r, min_col=2, min_row=2, max_row=len(ranked) + 1)
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(cats)
        chart.legend = None
        # Rows are written best-first; flip the category axis so rank 1 is at
        # the top of the chart, and keep the value axis at the bottom.
        chart.x_axis.scaling.orientation = "maxMin"
        chart.y_axis.crosses = "max"
        chart.y_axis.scaling.min = 0
        chart.y_axis.scaling.max = 4
        chart.dataLabels = DataLabelList(showVal=True)
        chart.width = 22
        chart.height = max(8.0, 0.55 * len(ranked) + 2)
        ws_r.add_chart(chart, "E2")

    # --- Candidate Feedback sheet ------------------------------------------
    with_feedback = [(r, ev) for r, ev in ok if ev.candidate_feedback]
    if with_feedback:
        ws_f = wb.create_sheet("Candidate Feedback")
        ws_f.append(["Name", "Email", "Feedback"])
        for _, ev in with_feedback:
            ws_f.append([ev.candidate_name, ev.email, ev.candidate_feedback])
            ws_f.cell(row=ws_f.max_row, column=3).alignment = wrap
        style_header(ws_f)
        for col, width in zip("ABC", [24, 32, 120]):
            ws_f.column_dimensions[col].width = width

    # --- Errors sheet ------------------------------------------------------
    failed = [r for r in results if r.error is not None]
    if failed:
        ws_e = wb.create_sheet("Errors")
        ws_e.append(["File", "Error"])
        for r in failed:
            ws_e.append([r.file.name, r.error])
        style_header(ws_e)
        ws_e.column_dimensions["A"].width = 30
        ws_e.column_dimensions["B"].width = 100

    wb.save(output)


# --------------------------------------------------------------------------
# CLI
# --------------------------------------------------------------------------

def read_arg_or_file(value: Optional[str], file_value: Optional[str], name: str, required: bool) -> str:
    if value and file_value:
        sys.exit(f"error: pass --{name} or --{name}-file, not both")
    if file_value:
        return Path(file_value).read_text(encoding="utf-8")
    if value:
        return value
    if required:
        sys.exit(f"error: --{name} or --{name}-file is required")
    return ""


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Evaluate resumes against skills and a project description; export to Excel."
    )
    parser.add_argument("resumes_dir", help="Directory containing resumes (.pdf, .docx, .txt, .md)")
    parser.add_argument("--prompt", help="Additional evaluation instructions")
    parser.add_argument("--prompt-file", help="File containing additional evaluation instructions")
    parser.add_argument("--project", help="Project description text")
    parser.add_argument(
        "--project-file",
        help=f"File containing the project description (default: {DEFAULT_PROJECT_FILE} if present)",
    )
    parser.add_argument("--skills", help="Comma-separated list of skills to evaluate (all weight 1.0)")
    parser.add_argument(
        "--skills-file",
        help=f"CSV file with `skill[,weight]` per line; weight defaults to 1.0 "
        f"(default: {DEFAULT_SKILLS_FILE}, or legacy {LEGACY_SKILLS_FILE}, if present)",
    )
    parser.add_argument("--output", default="evaluations.xlsx", help="Output .xlsx path (default: evaluations.xlsx)")
    parser.add_argument(
        "--provider",
        default=os.environ.get("PROVIDER", "ica"),
        choices=["ica", "anthropic"],
        help="LLM backend (default: ica — IBM Consulting Advantage)",
    )
    parser.add_argument("--host", help="Base URL override for the ICA backend (the /ica root)")
    parser.add_argument(
        "--model",
        default=None,
        help=f"Claude model ID (default: env ICA_MODEL / CLAUDE_MODEL, else {DEFAULT_MODEL})",
    )
    parser.add_argument("--workers", type=int, default=4, help="Concurrent evaluations (default: 4)")
    parser.add_argument(
        "--team-size",
        type=int,
        help="Target team size: adds a Team tab with the strongest mix of N candidates "
        "and per-skill score totals (or four cost-shaped teams when rates are available)",
    )
    parser.add_argument(
        "--rates",
        help="Rates table .xlsx with columns Country, Line, Band, Rate, Currency "
        f"(default: first of {', '.join(DEFAULT_RATES_CANDIDATES)} that exists)",
    )
    parser.add_argument(
        "--staff-rates",
        default=DEFAULT_STAFF_RATES_FILE,
        help=f"Staff rates workbook to create/update (default: {DEFAULT_STAFF_RATES_FILE}; "
        "pass an empty string to disable)",
    )
    parser.add_argument(
        "--screen",
        action="store_true",
        help="Two-tier mode: screen every resume with a cheaper model first; only candidates "
        "at or above --screen-cutoff get the full evaluation with --model",
    )
    parser.add_argument(
        "--screen-model",
        default=DEFAULT_SCREEN_MODEL,
        help=f"Model for the screening pass (default: {DEFAULT_SCREEN_MODEL})",
    )
    parser.add_argument(
        "--screen-cutoff",
        type=int,
        default=2,
        help="Minimum screening score (0-4) to advance to full evaluation (default: 2)",
    )
    args = parser.parse_args()

    resumes_dir = Path(args.resumes_dir)
    if not resumes_dir.is_dir():
        sys.exit(f"error: {resumes_dir} is not a directory")

    project = read_arg_or_file(args.project, args.project_file, "project", required=False)
    if not project:
        if Path(DEFAULT_PROJECT_FILE).is_file():
            project = Path(DEFAULT_PROJECT_FILE).read_text(encoding="utf-8")
        else:
            sys.exit(
                f"error: provide --project/--project-file, or create {DEFAULT_PROJECT_FILE} "
                "in the current directory"
            )

    skills_text = read_arg_or_file(args.skills, args.skills_file, "skills", required=False)
    if skills_text:
        if args.skills_file:
            skills = parse_skills_csv(skills_text, args.skills_file)
        else:
            skills = [Skill(name=s.strip()) for s in skills_text.split(",") if s.strip()]
    else:
        for candidate_file in (DEFAULT_SKILLS_FILE, LEGACY_SKILLS_FILE):
            path = Path(candidate_file)
            if path.is_file():
                skills = parse_skills_csv(path.read_text(encoding="utf-8"), candidate_file)
                break
        else:
            sys.exit(
                f"error: provide --skills/--skills-file, or create {DEFAULT_SKILLS_FILE} "
                "in the current directory"
            )
    if not skills:
        sys.exit("error: no skills provided")

    prompt = read_arg_or_file(args.prompt, args.prompt_file, "prompt", required=False)

    files = sorted(
        p for p in resumes_dir.iterdir()
        if p.is_file() and p.suffix.lower() in RESUME_EXTENSIONS and not p.name.startswith(".")
    )
    if not files:
        sys.exit(f"error: no resumes ({', '.join(sorted(RESUME_EXTENSIONS))}) found in {resumes_dir}")

    model = resolve_model(args.provider, args.model)
    client = make_client(args.provider, args.host)

    # Bootstrap: if staff-rates.xlsx doesn't exist yet, offer to populate it
    # with candidate names only (cheap model, no skill evaluations) so the
    # staffing team can fill in Geo/Line/Band before the full run.
    staff_path = Path(args.staff_rates) if args.staff_rates else None
    if staff_path is not None and not staff_path.is_file() and sys.stdin.isatty():
        answer = input(
            f"{staff_path} does not exist. Populate it with candidate names only, "
            "skipping skill evaluations? [y/N] "
        ).strip().lower()
        if answer in ("y", "yes"):
            print(f"Extracting candidate names with {args.screen_model}...")
            names: List[str] = []
            with ThreadPoolExecutor(max_workers=max(1, args.workers)) as name_pool:
                futures = {
                    name_pool.submit(extract_name, client, args.screen_model, p): p
                    for p in files
                }
                for future in as_completed(futures):
                    path = futures[future]
                    try:
                        name = future.result()
                        names.append(name)
                        print(f"  {path.name}: {name}")
                    except Exception as exc:
                        print(f"  {path.name}: FAILED ({type(exc).__name__}: {exc})")
            sync_staff_rates(staff_path, sorted(names), find_rates_table(args.rates))
            print(
                f"\nWrote {staff_path} with {len(names)} candidate(s). Fill in "
                "Geo/Line/Band (dropdowns provided) and re-run for rates and evaluations."
            )
            return 0

    print(
        f"Evaluating {len(files)} resume(s) against {len(skills)} skill(s) "
        f"with {model} via {args.provider}..."
    )
    system_prompt = build_system_prompt(prompt, project, skills)

    lock = threading.Lock()

    def run_all(paths: List[Path], worker, label: str) -> List[Result]:
        out: List[Result] = []
        done = 0
        with ThreadPoolExecutor(max_workers=max(1, args.workers)) as pool:
            futures = {pool.submit(worker, p): p for p in paths}
            for future in as_completed(futures):
                result = future.result()
                with lock:
                    out.append(result)
                    done += 1
                    status = "ok" if result.evaluation else f"FAILED ({result.error})"
                    print(f"  [{label} {done}/{len(paths)}] {result.file.name}: {status}")
        return out

    def full_eval(path: Path) -> Result:
        try:
            ev = evaluate_resume(client, model, system_prompt, path)
            return Result(file=path, evaluation=ev)
        except Exception as exc:  # keep going; failures land on the Errors sheet
            return Result(file=path, error=f"{type(exc).__name__}: {exc}")

    if args.screen:
        def screen_one(path: Path) -> Result:
            try:
                sr = screen_resume(client, args.screen_model, system_prompt, path)
                ev = ResumeEvaluation(
                    candidate_name=sr.candidate_name,
                    overall_score=sr.overall_score,
                    recommendation=sr.recommendation,
                    summary=f"[Screened by {args.screen_model}] {sr.summary}",
                    strengths=[],
                    gaps=[],
                    skill_evaluations=[],
                    email="",
                    candidate_feedback="",
                )
                return Result(file=path, evaluation=ev, stage="screened out")
            except Exception as exc:
                return Result(file=path, error=f"{type(exc).__name__}: {exc}")

        print(f"Screening pass with {args.screen_model} (cutoff: {args.screen_cutoff})...")
        screened = run_all(files, screen_one, "screen")
        advancing = [
            r.file for r in screened
            if r.evaluation and r.evaluation.overall_score >= args.screen_cutoff
        ]
        results = [
            r for r in screened
            if r.error or (r.evaluation and r.evaluation.overall_score < args.screen_cutoff)
        ]
        print(
            f"Screening done: {len(advancing)} of {len(files)} advance to full evaluation "
            f"({len(files) - len(advancing)} screened out or failed)."
        )
        if advancing:
            print(f"Full evaluation with {model}...")
            results += run_all(advancing, full_eval, "eval")
    else:
        results = run_all(files, full_eval, "eval")

    # --- Staff rates sync --------------------------------------------------
    staff = None
    evaluated_names = [
        r.evaluation.candidate_name
        for r in results
        if r.evaluation and r.stage == "full" and r.evaluation.skill_evaluations
    ]
    if args.staff_rates and evaluated_names:
        staff = sync_staff_rates(
            Path(args.staff_rates), evaluated_names, find_rates_table(args.rates)
        )

    # --- Team selection / team shapes --------------------------------------
    team = None
    shapes = None
    if args.team_size:
        pool = [
            (r, r.evaluation)
            for r in results
            if r.evaluation and r.stage == "full" and r.evaluation.skill_evaluations
        ]
        if not pool:
            print("warning: --team-size given but no fully-evaluated candidates; skipping Team sheet")
        else:
            size = min(args.team_size, len(pool))
            if size < args.team_size:
                print(
                    f"warning: only {len(pool)} fully-evaluated candidate(s); "
                    f"team size reduced to {size}"
                )

            def staff_info(ev) -> Optional[StaffInfo]:
                return (staff or {}).get(ev.candidate_name.strip().lower())

            def rate_of(entry: tuple) -> float:
                info = staff_info(entry[1])
                return info.rate if info is not None and info.rate is not None else 0.0

            missing_rates = []
            for _, ev in pool:
                info = staff_info(ev)
                if info is None or info.rate is None:
                    missing_rates.append(ev.candidate_name)
            if staff is not None and not missing_rates:
                # Every evaluated candidate has a cost rate → four team shapes.
                by_name = {ev.candidate_name.strip().lower(): (r, ev) for r, ev in pool}
                print(f"All candidates have rates — selecting four team shapes of {size} with {model}...")

                def resolve_members(names: List[str]) -> List[tuple]:
                    members: List[tuple] = []
                    for name in names:
                        entry = by_name.get(name.strip().lower())
                        if entry is not None and entry not in members:
                            members.append(entry)
                    if len(members) < size:
                        for entry in sorted(pool, key=lambda p: p[1].overall_score, reverse=True):
                            if entry not in members:
                                members.append(entry)
                            if len(members) == size:
                                break
                    return members[:size]

                try:
                    sel = select_team_shapes(client, model, system_prompt, pool, staff, size)
                    returned = {s.shape: s for s in sel.shapes}
                    shapes = []
                    for shape_name in TEAM_SHAPES:
                        entry = returned.get(shape_name)
                        members = resolve_members(entry.members if entry else [])
                        if entry:
                            rationale = entry.rationale
                            if entry.composition_notes:
                                rationale += f"\nComposition: {entry.composition_notes}"
                        else:
                            rationale = "(model omitted this shape; filled by overall score)"
                        shapes.append(shape_metrics(shape_name, members, rationale, skills, staff))
                except Exception as exc:
                    print(f"warning: team-shape selection failed ({type(exc).__name__}: {exc}); "
                          "using deterministic fallback shapes")
                    by_rate = sorted(pool, key=rate_of)
                    mid = max(0, (len(by_rate) - size) // 2)
                    fallback = {
                        "low cost": by_rate[:size],
                        "medium cost": by_rate[mid:mid + size],
                        "high cost": by_rate[-size:],
                        "cost not considered": greedy_best_fit(pool, skills, size),
                    }
                    shapes = [
                        shape_metrics(name, members, "Fallback: selected by cost rate ordering "
                                      "(model-based selection failed).", skills, staff)
                        for name, members in fallback.items()
                    ]
            elif staff is not None and missing_rates:
                print(
                    "note: no team shapes — missing cost rates for: "
                    + ", ".join(missing_rates)
                    + f"; populate Geo/Line/Band in {args.staff_rates} and re-run. "
                    "Building the single Team tab instead."
                )

            if shapes is None:
                print(f"Selecting the strongest team of {size} with {model}...")
                try:
                    selection = select_team(client, model, system_prompt, pool, size)
                    by_name = {ev.candidate_name.strip().lower(): (r, ev) for r, ev in pool}
                    members = []
                    for name in selection.selected_candidates:
                        entry = by_name.get(name.strip().lower())
                        if entry is not None and entry not in members:
                            members.append(entry)
                    # Top up with best-by-overall if the model returned unknown
                    # names or too few; trim if too many.
                    if len(members) < size:
                        for entry in sorted(pool, key=lambda p: p[1].overall_score, reverse=True):
                            if entry not in members:
                                members.append(entry)
                            if len(members) == size:
                                break
                    rationale = selection.rationale
                    if selection.composition_notes:
                        rationale += f"\nComposition: {selection.composition_notes}"
                    team = (members[:size], rationale)
                except Exception as exc:
                    print(f"warning: team selection failed ({type(exc).__name__}: {exc}); "
                          f"falling back to top {size} by overall score")
                    members = sorted(pool, key=lambda p: p[1].overall_score, reverse=True)[:size]
                    team = (members, f"Fallback: top {size} candidates by overall score "
                                     "(model-based team selection failed).")

    output = Path(args.output)
    write_workbook(results, skills, output, team=team, shapes=shapes, staff=staff)

    succeeded = sum(1 for r in results if r.evaluation)
    print(f"\nWrote {output} — {succeeded} evaluated, {len(results) - succeeded} failed.")
    return 0 if succeeded else 1


if __name__ == "__main__":
    sys.exit(main())
